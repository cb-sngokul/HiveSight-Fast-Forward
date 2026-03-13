#!/usr/bin/env python3
"""Generate HiveSight Architecture Document as DOCX."""

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('HiveSight', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    subtitle = doc.add_paragraph('Technical Architecture & Design Document')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # 1. Introduction
    add_heading(doc, '1. Introduction', 1)
    add_para(doc, 'HiveSight is a Temporal Billing Validation Engine for Chargebee subscriptions. It simulates subscription lifecycle over a user-provided period (start_month and end_month), including scheduled changes (Ramps), and validates termination dates. The simulation window is configurable—users specify the date range; there is no fixed 18-month limit. The application helps catch "time-bomb" errors before they occur—particularly around scheduled changes that the standard Chargebee UI may not fully reflect.')
    add_para(doc, 'Key capabilities:')
    for item in [
        'Simulate subscription billing over a configurable period',
        'Apply Chargebee Ramps (scheduled changes) during simulation',
        'Validate Ghost of March (cancellation date accuracy)',
        'AI-powered summary, alerts, and chat about simulation results',
        'Monthly breakdown with invoice trend visualization'
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2. Tech Stack
    add_heading(doc, '2. Technology Stack', 1)
    add_heading(doc, '2.1 Backend', 2)
    doc.add_paragraph('Java 21, Spring Boot 3.2.5, Spring Web (REST)')
    add_para(doc, 'The backend is a monolithic Spring Boot application exposing REST APIs. No Chargebee SDK is used; integration is via RestTemplate and Chargebee REST API v2.')
    add_heading(doc, '2.2 Frontend', 2)
    doc.add_paragraph('HTML5, CSS3, Bootstrap 5.3, Chart.js 4.4, Vanilla JavaScript')
    add_para(doc, 'Single-page application with no build step. All frontend logic is in app.js.')
    add_heading(doc, '2.3 External Integrations', 2)
    doc.add_paragraph('Chargebee REST API v2, LLM APIs (Groq, Grok, OpenAI, Gemini)')
    
    # 3. System Architecture
    add_heading(doc, '3. System Architecture', 1)
    add_para(doc, 'High-level architecture diagram:')
    add_code(doc, '''┌─────────────────────────────────────────────────────────────────────────────┐
│                              HiveSight                                       │
├─────────────────────────────────────────────────────────────────────────────┤
│  Frontend (index.html + app.js + hivesight.css)                             │
│  └─ Bootstrap 5, Chart.js, vanilla JS                                        │
├─────────────────────────────────────────────────────────────────────────────┤
│  ApiController (/api/*)              AiController (/api/ai/*)               │
│  └─ subscriptions, simulate,         └─ summarize, chat, alerts             │
│     validate, details, debug                                                 │
├──────────────────────────────────────┬──────────────────────────────────────┤
│  ChargebeeService                    │  AiService                            │
│  └─ Chargebee REST API               │  └─ Groq/Grok/OpenAI/Gemini           │
├──────────────────────────────────────┴──────────────────────────────────────┤
│  Simulator (engine)                                                          │
│  └─ SimulatedSubscription, Ramp, TimelineEvent, MonthlyBreakdown              │
│  └─ BillingPeriodUtil                                                        │
└─────────────────────────────────────────────────────────────────────────────┘''')
    
    # 4. Package Structure
    add_heading(doc, '4. Package Structure', 1)
    add_code(doc, '''com.hivesight/
├── HiveSightApplication.java          # Spring Boot entry point
├── controller/
│   ├── ApiController.java             # REST API (subscriptions, simulation, validation)
│   └── AiController.java              # AI endpoints (summarize, chat, alerts)
├── engine/
│   ├── Simulator.java                 # Core simulation engine
│   ├── BillingPeriodUtil.java         # Billing period arithmetic
│   ├── TimelineEvent.java             # Timeline event model
│   ├── MonthlyBreakdown.java          # Per-month invoice breakdown
│   ├── Ramp.java                      # Chargebee ramp model
│   └── SimulatedSubscription.java     # Subscription state for simulation
└── service/
    ├── ChargebeeService.java          # Chargebee API client
    └── AiService.java                 # LLM integration''')
    
    # 5. Core Components
    add_heading(doc, '5. Core Components', 1)
    
    add_heading(doc, '5.1 Simulator (Engine)', 2)
    add_para(doc, 'The Simulator is the heart of HiveSight. It iterates through time from simulationStart to effectiveEnd, applying ramps when due, handling trial end, pause/resume, cancellation, and contract term end. For each renewal, it emits TimelineEvents and builds MonthlyBreakdown records.')
    add_para(doc, 'Key methods:')
    doc.add_paragraph('simulate() - Main simulation loop', style='List Bullet')
    doc.add_paragraph('subscriptionEndDate() - Computes end from cancelled_at, cancel_at, or contract_end', style='List Bullet')
    doc.add_paragraph('toSimulatedSubscription() - Maps Chargebee JSON to SimulatedSubscription', style='List Bullet')
    doc.add_paragraph('buildMonthlyBreakdown() - Per-month invoice with changes (price, frequency, items)', style='List Bullet')
    doc.add_paragraph('formatBillingPeriod() - Human-readable (quarterly, monthly, yearly)', style='List Bullet')
    
    add_heading(doc, '5.2 ChargebeeService', 2)
    add_para(doc, 'REST client for Chargebee. Uses RestTemplate with Basic Auth. Base URL: https://{site}.chargebee.com/api/v2')
    add_para(doc, 'Key operations: listSubscriptions(), getSubscription(), listRamps(), getSubscriptionDetails(), simulate(), validateGhostOfMarch()')
    
    add_heading(doc, '5.3 AiService', 2)
    add_para(doc, 'LLM integration supporting Groq, Grok, OpenAI, and Gemini. Converts simulation data to AI-friendly format (cents→dollars), builds structured prompts, and post-processes outputs (strip preambles, fix amount display).')
    add_para(doc, 'Features: generateSummary(), generateAlerts(), chat(). Provider selected via ai.provider config.')
    
    # 6. Data Models
    add_heading(doc, '6. Data Models', 1)
    
    add_heading(doc, '6.1 SimulationResult', 2)
    add_code(doc, '''record SimulationResult(
    String subscriptionId, String customerId,
    long simulationStart, long simulationEnd, Long subscriptionEndDate,
    List<TimelineEvent> events, List<MonthlyBreakdown> monthlyBreakdowns,
    Long chargebeeUiNextBilling, Long hivesightNextBilling,
    String currencyCode, String timezone,
    Boolean validationPassed, String validationMessage)''')
    
    add_heading(doc, '6.2 TimelineEvent', 2)
    add_para(doc, 'Types: renewal, ramp_applied, cancelled, trial_end, paused, resumed, non_renewing, contract_end')
    add_code(doc, 'record TimelineEvent(String type, long date, String dateFormatted, String description, Ramp ramp, Long amount, String currencyCode)')
    
    add_heading(doc, '6.3 MonthlyBreakdown', 2)
    add_code(doc, '''record MonthlyBreakdown(
    String monthKey, String monthLabel, long unitPrice, int quantity,
    long subtotalCents, long discountCents, Integer taxRatePercent, long taxCents, long totalCents,
    String currencyCode, List<String> changes, Long impactVsPreviousCents)''')
    add_para(doc, 'changes: "Price changed from X to Y", "Plan frequency changed from quarterly to monthly", "Item(s) added", etc.')
    
    add_heading(doc, '6.4 Ramp', 2)
    add_para(doc, 'Chargebee ramp model with itemsToAdd, itemsToUpdate, itemsToRemove, discountsToAdd, etc.')
    
    add_heading(doc, '6.5 SimulatedSubscription', 2)
    add_para(doc, 'Subscription state for simulation: billing period, items, cancelledAt, contract term, trial, pause/resume.')
    
    # 7. API Design
    add_heading(doc, '7. API Design', 1)
    
    add_heading(doc, '7.1 ApiController (/api)', 2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Method'
    hdr[1].text = 'Endpoint'
    hdr[2].text = 'Params'
    hdr[3].text = 'Description'
    for row in [
        ('GET', '/health', '-', 'Health check'),
        ('GET', '/subscriptions', 'has_scheduled_changes', 'List subscriptions'),
        ('GET', '/simulate/{id}', 'start_month, end_month, timezone', 'Run simulation'),
        ('GET', '/subscription/{id}/details', '-', 'Subscription details card'),
        ('GET', '/subscription/{id}/debug', '-', 'Debug cancellation fields'),
        ('GET', '/validate/ghost-of-march/{id}', 'expected_cancel, start_month, end_month, timezone', 'Validate cancellation date'),
    ]:
        r = table.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = row
    
    add_heading(doc, '7.2 AiController (/api/ai)', 2)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr2 = table2.rows[0].cells
    hdr2[0].text, hdr2[1].text, hdr2[2].text, hdr2[3].text = 'Method', 'Endpoint', 'Body', 'Description'
    for row in [
        ('GET', '/status', '-', 'AI enabled/disabled'),
        ('POST', '/summarize', 'data', 'Generate AI summary'),
        ('POST', '/chat', 'message, data', 'Chat about simulation'),
        ('POST', '/alerts', 'data', 'Generate risk alerts'),
    ]:
        r = table2.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = row
    
    # 8. Key Flows
    add_heading(doc, '8. Key Flows', 1)
    
    add_heading(doc, '8.1 Simulation Flow', 2)
    add_para(doc, 'The simulation period is user-provided via start_month and end_month (YYYY-MM format). If omitted, defaults apply (current month to ~18 months ahead).')
    doc.add_paragraph('1. Client calls GET /api/simulate/{id}?start_month=2026-01&end_month=2027-12', style='List Number')
    doc.add_paragraph('2. ChargebeeService fetches subscription + ramps from Chargebee', style='List Number')
    doc.add_paragraph('3. Simulator.toSimulatedSubscription() maps to internal model', style='List Number')
    doc.add_paragraph('4. Simulator.subscriptionEndDate() from cancelled_at, cancel_at, or contract_end', style='List Number')
    doc.add_paragraph('5. Simulator.simulate() loops, applies ramps, emits events, builds monthly breakdowns', style='List Number')
    doc.add_paragraph('6. Returns SimulationResult (JSON)', style='List Number')
    
    add_heading(doc, '8.2 AI Summary Flow', 2)
    doc.add_paragraph('1. Client POSTs { data: simulationResult } to /api/ai/summarize', style='List Number')
    doc.add_paragraph('2. AiService.formatAmountsForAi() converts cents to dollars, adds simulationWindowDisplay, subscriptionBehavior', style='List Number')
    doc.add_paragraph('3. buildSummaryFields() extracts renewalsCount, totalProjectedRevenueDisplay, rampsCount, notableChanges, priceRange', style='List Number')
    doc.add_paragraph('4. Prompt instructs LLM to use exact values only', style='List Number')
    doc.add_paragraph('5. stripSummaryPreamble() cleans output', style='List Number')
    
    add_heading(doc, '8.3 Ghost of March Validation', 2)
    add_para(doc, 'Validates that a subscription scheduled to cancel actually terminates on the expected date. Compares last event (cancelled) date with expected_cancel parameter.')
    
    # 9. Configuration
    add_heading(doc, '9. Configuration', 1)
    add_para(doc, 'application.properties:')
    add_code(doc, '''chargebee.site, chargebee.api-key, chargebee.timezone, chargebee.default-tax-rate
ai.provider (groq|grok|openai|gemini)
ai.groq.api-key, ai.groq.model
ai.grok.api-key, ai.grok.model
ai.openai.api-key, ai.openai.model
ai.gemini.api-key, ai.gemini.model
server.port (default 8766)''')
    
    # 10. Frontend Architecture
    add_heading(doc, '10. Frontend Architecture', 1)
    add_para(doc, 'Single-page app. Key modules in app.js:')
    doc.add_paragraph('renderTimeline() - Event timeline with icons', style='List Bullet')
    doc.add_paragraph('renderMonthlyBreakdown() - Per-month cards with unit price, subtotal, tax, total, impact', style='List Bullet')
    doc.add_paragraph('buildInvoiceTrendData() / renderInvoiceTrendChart() - Chart.js line chart', style='List Bullet')
    doc.add_paragraph('renderDetailedReport() - Stats (renewals, ramps, window, revenue, cancelled)', style='List Bullet')
    doc.add_paragraph('callAi() - POST to /api/ai/* with lastSimulationData', style='List Bullet')
    doc.add_paragraph('fixReportAmounts() - Post-process AI output (correct $15,000→$150, auto-renew phrasing)', style='List Bullet')
    
    # 11. Security & Deployment
    add_heading(doc, '11. Security & Deployment', 1)
    add_para(doc, 'API keys in application.properties; recommend env vars for production. CORS: CrossOrigin(origins="*") on controllers. No authentication layer—suitable for internal/demo use.')
    add_para(doc, 'Build: ./mvnw clean package -DskipTests')
    add_para(doc, 'Run: java -jar target/hivesight-1.0.0.jar')
    
    # Save
    out_path = '/Users/gokulrajsn/Desktop/Hivesight/docs/HiveSight_Architecture_Design.docx'
    import os
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"Created: {out_path}")

if __name__ == '__main__':
    main()
